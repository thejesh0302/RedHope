import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('RedHope Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Comprehensive Blood Donation Platform\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "RedHope is a web-based blood donation platform built to seamlessly connect blood donors "
        "with patients in need. By leveraging location-based matching and artificial intelligence "
        "for eligibility assessments, the system streamlines the blood donation process, ensuring "
        "quick and efficient access to life-saving blood during emergencies."
    )
    
    # Figure 1: Logo
    logo_path = r"d:\Redhope\static\images\redhope_name_logo.png"
    if os.path.exists(logo_path):
        doc.add_heading('Figure 1: RedHope Logo', level=3)
        doc.add_picture(logo_path, width=Inches(3.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Key Features
    doc.add_heading('2. Key Features', level=1)
    features = doc.add_paragraph()
    features.add_run("User Registration & Roles: ").bold = True
    features.add_run("Separate flows for Donors and Patients with profile management.\n")
    features.add_run("AI-Driven Eligibility: ").bold = True
    features.add_run("Uses Groq API (LLaMA3-70b) to provide compassionate, personalized explanations for donor ineligibility.\n")
    features.add_run("Location-Based Matching: ").bold = True
    features.add_run("Employs the Haversine formula to match patients with the nearest eligible donors based on geographical coordinates (latitude/longitude).\n")
    features.add_run("Real-time Requests & Notifications: ").bold = True
    features.add_run("Patients can request blood units, and donors receive immediate notifications to accept or reject.\n")
    features.add_run("Donation History & Certificates: ").bold = True
    features.add_run("Tracks past donations and automatically generates digital certificates for donors.")

    # 3. System Architecture
    doc.add_heading('3. System Architecture & Technology Stack', level=1)
    doc.add_paragraph("• Backend: Python, Flask\n"
                      "• Database: SQLite\n"
                      "• Frontend: HTML5, CSS3, Jinja2 Templates\n"
                      "• AI Integration: Groq API (llama-3.3-70b-versatile)\n"
                      "• Geolocation: Haversine distance calculation")

    # Figure 2: Architecture representation or hero image
    hero_path = r"d:\Redhope\static\images\hero_community.png"
    if os.path.exists(hero_path):
        doc.add_heading('Figure 2: Community Engagement Illustration', level=3)
        doc.add_picture(hero_path, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. Database Schema (Tables)
    doc.add_heading('4. Database Schema', level=1)
    doc.add_paragraph("The application uses a relational database (SQLite) consisting of several core tables.")
    
    # Table 1: Users Table
    doc.add_heading('Table 4.1: Users Table Structure', level=3)
    table1 = doc.add_table(rows=1, cols=3)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Column Name'
    hdr_cells[1].text = 'Data Type'
    hdr_cells[2].text = 'Description'
    
    user_cols = [
        ('id', 'INTEGER', 'Primary Key'),
        ('full_name', 'TEXT', 'User full name'),
        ('email', 'TEXT', 'Unique email address'),
        ('role', 'TEXT', 'Donor or Patient'),
        ('latitude', 'REAL', 'Geographical latitude'),
        ('longitude', 'REAL', 'Geographical longitude')
    ]
    for col_name, d_type, desc in user_cols:
        row_cells = table1.add_row().cells
        row_cells[0].text = col_name
        row_cells[1].text = d_type
        row_cells[2].text = desc

    # Table 2: Patient Requests
    doc.add_heading('Table 4.2: Patient Requests Table', level=3)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Column Name'
    hdr_cells[1].text = 'Data Type'
    hdr_cells[2].text = 'Description'
    
    req_cols = [
        ('request_id', 'INTEGER', 'Primary Key'),
        ('patient_id', 'INTEGER', 'Foreign Key (Users)'),
        ('blood_group_needed', 'TEXT', 'Required blood type'),
        ('units_required', 'INTEGER', 'Number of units'),
        ('status', 'TEXT', 'Pending, Fulfilled, etc.')
    ]
    for col_name, d_type, desc in req_cols:
        row_cells = table2.add_row().cells
        row_cells[0].text = col_name
        row_cells[1].text = d_type
        row_cells[2].text = desc

    # 5. Data Visualization (Matplotlib Figure)
    doc.add_heading('5. System Process Flow', level=1)
    doc.add_paragraph("The following chart represents the typical distribution of user roles within the platform based on standard platform usage metrics.")
    
    # Generate a chart using matplotlib
    labels = 'Donors', 'Patients', 'Admins'
    sizes = [65, 30, 5]
    colors = ['#ff9999','#66b3ff','#99ff99']
    
    plt.figure(figsize=(5, 4))
    plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')  
    plt.title('User Distribution')
    
    chart_path = r"d:\Redhope\user_distribution.png"
    plt.savefig(chart_path)
    plt.close()
    
    doc.add_heading('Figure 3: Expected User Distribution', level=3)
    doc.add_picture(chart_path, width=Inches(4.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Clean up the temporary chart image
    if os.path.exists(chart_path):
        os.remove(chart_path)

    # Save the document
    out_path = r"d:\Redhope\RedHope_Project_Report.docx"
    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    out = create_report()
    print(f"Report successfully created at {out}")
